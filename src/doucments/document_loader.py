import os, re, json
from PIL import Image
import platform,chardet
import pdfplumber, pytesseract
from pdf2image import convert_from_path
from docx import Document as DocxDocument
from typing import List, Dict, Any, Optional
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
class DocumentProcessor:
    """
    文档处理主类（单文件整合版）
    包含三大核心组件：
    1. _OcrService: 负责图片识别与表格结构还原
    2. _PdfHandler: 负责PDF（扫描件/原生）的提取
    3. _WordHandler: 负责Word文档的提取
    """

    def __init__(
            self,
            chunk_size: int = 400,
            chunk_overlap: int = 180,
            enable_ocr: bool = True
    ):
        # 中文优化的语义切分器
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=[
                "\n\n# ", "\n## ", "\n### ",
                "\n\n", "\r\n\r\n",
                "。", "！", "？", "；", "：",
                ". ", "! ", "? ", "; ", ": ",
                "\n", "\r\n", " ", "\t", "",
            ],
        )
        self.enable_ocr = enable_ocr
        self._global_chunk_counter = 0

        # 层级匹配模式
        self._hierarchy_patterns = [
            (r"^(#{1,6})\s+(.+)$", "header"),
            (r"^([一二三四五六七八九十]+[、．.])\s*(.+)$", "chinese_num"),
            (r"^([（(][一二三四五六七八九十]+[）)])\s*(.+)$", "chinese_bracket"),
            (r"^(\d+[、．.])\s*(.+)$", "num"),
            (r"^([（(]\d+[）)])\s*(.+)$", "num_bracket"),
            (r"^([①②③④⑤⑥⑦⑧⑨⑩⑪⑫⑬⑭⑮⑯⑰⑱⑲⑳])\s*(.+)$", "circle_num"),
        ]

        # 初始化服务
        self.ocr_service = self._OcrService()
        self.pdf_handler = self._PdfHandler(self.ocr_service, self._get_poppler_path())
        self.word_handler = self._WordHandler(self.ocr_service)

        # 配置 Tesseract 路径
        self._setup_tesseract_path()

    # ==========================================
    # 辅助方法：层级检测与智能切分
    # ==========================================
    def _detect_hierarchy(self, text: str) -> List[Dict[str, Any]]:
        """检测文本中的层级结构"""
        hierarchy_info = []
        lines = text.split("\n")
        for line_num, line in enumerate(lines):
            line_stripped = line.strip()
            if not line_stripped:
                continue
            for pattern, pattern_type in self._hierarchy_patterns:
                match = re.match(pattern, line_stripped)
                if match:
                    level = len(match.group(1)) if pattern_type == "header" else 1
                    content = match.group(2).strip()
                    hierarchy_info.append({
                        "line_num": line_num,
                        "type": pattern_type,
                        "level": level,
                        "prefix": match.group(1),
                        "content": content
                    })
                    break
        return hierarchy_info

    def _smart_split_by_hierarchy(self, text: str, max_split_level: int = 2) -> List[str]:
        """
        根据层级进行粗切分
        :param text: 原始文本
        :param max_split_level: 最大允许切分的标题层级
        """
        hierarchy_info = self._detect_hierarchy(text)
        if not hierarchy_info:
            return [text]

        lines = text.split("\n")
        chunks = []
        current_start = 0
        # 思路：如果连续多个"标题"之间没有实质内容（行数很少），
        # 说明它们是目录条目，不作为切分点
        split_points = []
        i = 0
        while i < len(hierarchy_info):
            info = hierarchy_info[i]
            level = info.get("level", 99)
            end_line = info["line_num"]
            if level <= max_split_level and end_line > current_start:
                # 检查这个标题到下一个标题之间是否有足够的正文内容
                # 如果间距太小（<=2行），大概率是目录条目
                next_line = hierarchy_info[i + 1]["line_num"] if i + 1 < len(hierarchy_info) else len(lines)
                content_lines = next_line - end_line

                if content_lines > 2:
                    # 有实质内容，作为切分点
                    split_points.append(info)
                # 否则跳过，认为是目录条目
            i += 1

        # 如果没有有效的切分点，返回全文
        if not split_points:
            return [text]
        for info in split_points:
            end_line = info["line_num"]
            if end_line > current_start:
                chunk_text = "\n".join(lines[current_start:end_line]).strip()
                if chunk_text:
                    chunks.append(chunk_text)
                current_start = end_line
        # 添加最后一块（从最后一个切分点到文末）—— 在循环外部
        if current_start < len(lines):
            last_chunk = "\n".join(lines[current_start:]).strip()
            if last_chunk:
                chunks.append(last_chunk)

        return chunks if chunks else [text]

    # ==========================================
    # 内部类 1: OCR 服务
    # ==========================================
    class _OcrService:
        """负责图片转文字及表格识别"""
        def recognize_image(self, image) -> str:
            try:
                # 【新增】类型检查和转换，兼容 str, bytes, PIL.Image
                if isinstance(image, str):
                    # 如果传入的是文件路径，则打开它
                    image = Image.open(image)
                elif isinstance(image, bytes):
                    # 如果是字节流，尝试从字节流加载
                    from io import BytesIO
                    image = Image.open(BytesIO(image))
                elif not isinstance(image, Image.Image):
                    # 其他未知类型
                    raise ValueError(f"Unsupported image type: {type(image)}")
                # 确保图片已加载到内存，防止源文件被占用或关闭
                image.load()
                text = pytesseract.image_to_string(image, lang="chi_sim+eng")
                return text.strip()
            except Exception as e:
                print(f"OCR 识别失败: {e}")
                return ""

    # ==========================================
    # 内部类 2: PDF 处理器
    # ==========================================
    class _PdfHandler:
        def __init__(self, ocr_service: "DocumentProcessor._OcrService", poppler_path: Optional[str]):
            self.ocr_service = ocr_service
            self.poppler_path = poppler_path

        def extract_text(self, file_path: str) -> str:
            full_text = ""
            # 【修改】禁用目录提取，避免污染正文
            try:
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        text = page.extract_text() or ""
                        tables = page.extract_tables() or []

                        # 拼接表格内容
                        table_texts = []
                        for table in tables:
                            rows = [" | ".join([str(cell) if cell else "" for cell in row]) for row in table]
                            table_texts.append("\n".join(rows))

                        page_content = text + "\n\n" + "\n\n".join(table_texts)

                        # 如果提取出的纯文本太少，触发 OCR
                        if len(text.strip()) < 50 and self.ocr_service:
                            images = convert_from_path(
                                file_path,
                                first_page=page.page_number,
                                last_page=page.page_number,
                                poppler_path=self.poppler_path
                            )
                            if images:
                                ocr_text = self.ocr_service.recognize_image(images[0])
                                page_content = ocr_text

                        full_text += page_content + "\n\n---PAGE BREAK---\n\n"
            except Exception as e:
                print(f"pdfplumber 解析失败，尝试全量 OCR: {e}")
                # 降级方案...
            return full_text.strip()
        def _extract_outline(self, file_path: str) -> str:
            """
            PDF 目录提取（已禁用）原因：目录条目会被 _detect_hierarchy 重复识别为标题，
            """
            return ""

    # ==========================================
    # 内部类 3: Word 处理器
    # ==========================================
    class _WordHandler:
        def __init__(self, ocr_service: "DocumentProcessor._OcrService"):
            self.ocr_service = ocr_service

        def extract_text(self, file_path: str) -> str:
            doc = DocxDocument(file_path)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            # 提取图片并进行 OCR
            for rel in doc.part.rels.values():
                if "image" in rel.reltype:
                    try:
                        # 【修改】直接使用字节流，不再创建临时文件
                        image_blob = rel.target_part.blob
                        # 将字节流直接传递给 OCR 服务
                        ocr_text = self.ocr_service.recognize_image(image_blob)
                        if ocr_text:
                            full_text.append("\n[Image Content]:\n" + ocr_text)
                    except Exception as e:
                        print(f"Word 图片 OCR 失败: {e}")
            return "\n\n".join(full_text).strip()

    # ==========================================
    # 主类的辅助方法 (路径配置)
    # ==========================================
    def _setup_tesseract_path(self):
        """智能配置 Tesseract OCR 引擎路径"""
        if platform.system() != "Windows":
            return
        default_paths = [
            r"E:\damoxing\python-project\help_ocr\tesseract\install\tesseract.exe",
            r"C:\Program Files\Tesseract-OCR\tesseract.exe",
            r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
        ]
        for path in default_paths:
            if os.path.exists(path):
                pytesseract.pytesseract.tesseract_cmd = path
                print(f"自动发现 Tesseract 路径: {path}")
                return
        print("警告: 未在 Windows 上找到 Tesseract，OCR 功能将无法使用！")

    def _get_poppler_path(self):
        """智能获取 Poppler 路径"""
        if platform.system() != "Windows":
            return None
        default_paths = [
            r"E:\damoxing\python-project\help_ocr\poppler-26.02.0\Library\bin",
            r"C:\Program Files\poppler\Library\bin",
            r".\poppler\Library\bin",
        ]
        for path in default_paths:
            exe_name = "pdftoppm.exe"
            if os.path.exists(os.path.join(path, exe_name)):
                print(f"自动发现 Poppler 路径: {path}")
                return path
        print("警告: 未在 Windows 上找到 Poppler 路径，PDF 扫描件处理可能会失败。")
        return None

    # ==========================================
    # 主类的对外接口
    # ==========================================
    def load_document(self, file_path: str) -> List[Document]:
        ext = os.path.splitext(file_path)[1].lower()
        text = ""
        if ext == ".pdf":
            text = self.pdf_handler.extract_text(file_path)
            f_type = "pdf"
        elif ext == ".docx":
            text = self.word_handler.extract_text(file_path)
            f_type = "docx"
        elif ext == ".txt":
            with open(file_path, "rb") as f:
                raw = f.read(10240)
            result = chardet.detect(raw)
            encoding = result["encoding"] if result["confidence"] > 0.7 else "utf-8"
            try:
                with open(file_path, "r", encoding=encoding) as f:
                    text = f.read()
            except:
                with open(file_path, "r", encoding="gbk", errors="ignore") as f:
                    text = f.read()
            f_type = "txt"
        else:
            raise ValueError(f"不支持的格式: {ext}")

        if not text.strip():
            raise ValueError("文档内容为空")

        return [Document(page_content=text, metadata={"source": file_path, "file_type": f_type})]

    def load_directory(self, directory_path: str) -> List[Document]:
        all_documents = []
        try:
            for root, _, files in os.walk(directory_path):
                for file_name in files:
                    file_path = os.path.join(root, file_name)
                    ext = os.path.splitext(file_path)[1].lower()
                    if ext not in [".txt", ".pdf", ".docx"]:
                        continue
                    docs = self.load_document(file_path)
                    all_documents.extend(docs)
            return self.split_documents(all_documents)
        except Exception as e:
            print(f"加载目录失败: {e}")
            return []

    # ==================================================
    # 核心逻辑：文档切分
    # ==================================================
    def split_documents(self, documents: List[Document]) -> List[Document]:
        if self.text_splitter is None:
            raise ValueError("切分器未初始化")
        all_split_docs = []
        for doc in documents:
            content = doc.page_content
            lines = content.split("\n")
            hierarchy_info = self._detect_hierarchy(content)
            # 如果没有检测到层级，直接使用 LangChain 切分
            if not hierarchy_info:
                chunks = self.text_splitter.split_documents([doc])
                for chunk in chunks:
                    chunk.metadata["chunk_id"] = self._global_chunk_counter
                    self._global_chunk_counter += 1
                    all_split_docs.append(chunk)
                continue

            # 1. 提取所有标题行号和信息
            header_spans = []
            current_header = None
            for line_num, line in enumerate(lines):
                line_stripped = line.strip()
                if not line_stripped:
                    continue
                is_header = False
                header_data = None
                for pattern, pattern_type in self._hierarchy_patterns:
                    if re.match(pattern, line_stripped):
                        is_header = True
                        header_data = {
                            "line_num": line_num,
                            "type": pattern_type,
                            "level": len(line_stripped.split()[0]) if pattern_type == "header" else 1,
                            "prefix": line_stripped.split()[0],
                            "content": " ".join(line_stripped.split()[1:]),
                        }
                        break
                if is_header and header_data["level"] > 2:
                    is_header = False  # 强制取消它的切分点资格
                if is_header:
                    if current_header is not None and current_header["content"]:
                        header_spans.append({
                            "header": current_header,
                            "start": current_header["start_line"],
                            "end": line_num - 1
                        })
                    header_data["start_line"] = line_num
                    current_header = header_data

            if current_header is not None:
                header_spans.append({
                    "header": current_header,
                    "start": current_header["start_line"],
                    "end": len(lines) - 1
                })

            # 2. 根据切分出的块进行处理
            for span in header_spans:
                start = span["start"]
                end = span["end"]
                header_info = span["header"]
                block_lines = lines[start:end + 1]
                block_text = "\n".join(block_lines).strip()
                if not block_text:
                    continue
                temp_doc = Document(
                    page_content=block_text,
                    metadata=doc.metadata.copy()
                )
                chunks = self.text_splitter.split_documents([temp_doc])
                for chunk in chunks:
                    chunk.metadata["hierarchy"] = {
                        "type": header_info["type"],
                        "level": header_info["level"],
                        "title": header_info["content"]
                    }
                    title_prefix = "#" * header_info["level"] + " " + header_info["content"]
                    if not chunk.page_content.startswith(title_prefix):
                        chunk.page_content = title_prefix + "\n\n" + chunk.page_content
                    chunk.metadata["chunk_id"] = self._global_chunk_counter
                    self._global_chunk_counter += 1
                    all_split_docs.append(chunk)

        self.save_chunks_to_local(all_split_docs)
        # chromaDB 不支持 字典需要序列化为json字符串，
        # 需要将 None 转为空字符串
        clean_chunks = []
        for doc in all_split_docs:
            # 复制一份元数据，避免修改原对象
            new_meta = doc.metadata.copy()
            # 遍历所有元数据字段
            for k, v in new_meta.items():
                # 如果值是字典、列表或复杂对象，强制转为 JSON 字符串
                if isinstance(v, (dict, list)):
                    new_meta[k] = json.dumps(v, ensure_ascii=False)
                # ChromaDB 也不支持 None，建议转为空字符串
                elif v is None:
                    new_meta[k] = ""
            # 更新文档的元数据
            doc.metadata = new_meta
            clean_chunks.append(doc)

        return clean_chunks  # 返回清洗后的数据

    def save_chunks_to_local(self, documents: List[Document], output_dir: str = "data/chunk_documents") -> None:
        if not documents:
            print("警告: 没有文档块需要保存。")
            return
        os.makedirs(output_dir, exist_ok=True)
        # 使用 enumerate 获取全局索引，防止不同文件的 chunk_id 重复导致覆盖
        for global_idx, doc in enumerate(documents):
            chunk_id = doc.metadata.get("chunk_id", global_idx)
            source_path = doc.metadata.get("source", "unknown_file")
            # 处理文件名，确保唯一性
            base_name = os.path.splitext(os.path.basename(source_path))[0]
            # 格式示例: chunk_0001_policy_A.md
            filename = f"chunk_{global_idx:04d}_{base_name}.md"
            filepath = os.path.join(output_dir, filename)
            metadata_str = "---\n"
            for key, value in doc.metadata.items():
                value_str = str(value).replace("\n", " ")
                metadata_str += f"{key}: {value_str}\n"
            metadata_str += "---\n\n"
            content = metadata_str + doc.page_content
            try:
                with open(filepath, "w", encoding="utf-8") as f:
                    f.write(content)
            except Exception as e:
                print(f"保存文件失败 {filepath}: {e}")

    def process_documents(self, input_path: str) -> List[Document]:
        if os.path.isfile(input_path):
            return self.load_document(input_path)
        elif os.path.isdir(input_path):
            return self.load_directory(input_path)
        else:
            raise ValueError(f"无效路径: {input_path}")
# 全局实例
documentProcessor = DocumentProcessor()