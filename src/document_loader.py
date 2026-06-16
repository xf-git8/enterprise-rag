import os
import tempfile
import platform
import chardet
from typing import List

from langchain_core.documents import Document
from pdf2image import convert_from_path
from langchain_text_splitters import RecursiveCharacterTextSplitter
import pdfplumber
import pytesseract
from PIL import Image
from docx import Document as DocxDocument


class DocumentProcessor:
    """
    文档处理主类（单文件整合版）
    包含三大核心组件：
    1. _OcrService: 负责图片识别与表格结构还原
    2. _PdfHandler: 负责PDF（扫描件/原生）的提取
    3. _WordHandler: 负责Word文档的提取
    """

    def __init__(
        self, chunk_size: int = 500, chunk_overlap: int = 200, enable_ocr: bool = True
    ):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", " ", ""],
        )
        self.enable_ocr = enable_ocr
        self._global_chunk_counter = 0
        # 初始化服务
        self.ocr_service = self._OcrService()
        self.pdf_handler = self._PdfHandler(self.ocr_service, self._get_poppler_path())
        self.word_handler = self._WordHandler(self.ocr_service)
        # 配置 Tesseract 路径
        self._setup_tesseract_path()

    # ==========================================
    # 内部类 1: OCR 与 表格处理服务
    # ==========================================
    class _OcrService:
        """底层OCR服务，负责图片识别和表格格式化"""

        def __init__(self):
            pass

        def _detect_table_pattern(self, text: str) -> bool:
            """检测OCR结果是否为表格模式（简单启发式）"""
            lines = [line.strip() for line in text.split("\n") if line.strip()]
            if len(lines) < 2:
                return False
            # 如果多行包含竖线 | 或者 空格分隔的列数较一致
            pipe_lines = sum(1 for line in lines if "|" in line)
            return pipe_lines >= len(lines) * 0.5

        def _convert_to_markdown_table(self, text: str) -> str:
            """尝试将纯文本表格转换为Markdown格式"""
            lines = [line.strip() for line in text.split("\n") if line.strip()]
            if len(lines) < 2:
                return text

            rows = []
            for line in lines:
                if "|" in line:
                    cells = [c.strip() for c in line.split("|") if c.strip()]
                else:
                    cells = [c.strip() for c in line.split() if c.strip()]
                if len(cells) >= 2:
                    rows.append(cells)

            if len(rows) < 2:
                return text

            # 构建Markdown
            markdown = []
            # 表头
            markdown.append("| " + " | ".join(rows[0]) + " |")
            # 分隔线
            markdown.append("| " + " | ".join(["---"] * len(rows[0])) + " |")
            # 数据行
            for row in rows[1:]:
                while len(row) < len(rows[0]):
                    row.append("")
                markdown.append("| " + " | ".join(row) + " |")
            return "\n".join(markdown)

        def process_image(self, image: Image.Image) -> str:
            """处理单张图片，自动判断是表格还是文本"""
            try:
                # 配置参数
                config_table = "--psm 6"  # 假设为均匀表格
                config_normal = "--psm 3"  # 完全自动页面分割

                # 1. 尝试表格识别
                table_text = pytesseract.image_to_string(
                    image, lang="chi_sim+eng", config=config_table
                )
                if self._detect_table_pattern(table_text):
                    md_table = self._convert_to_markdown_table(table_text)
                    return f"[图片表格识别]:\n{md_table}\n"

                # 2. 普通文本识别
                normal_text = pytesseract.image_to_string(
                    image, lang="chi_sim+eng", config=config_normal
                )
                return f"[图片文字识别]:\n{normal_text}\n"

            except Exception as e:
                return f"[OCR识别错误]: {str(e)}\n"

    # ==========================================
    # 内部类 2: PDF 处理器
    # ==========================================
    class _PdfHandler:
        """专门处理PDF逻辑"""

        def __init__(self, ocr_service: "_OcrService", poppler_path: str = None):
            self.ocr_service = ocr_service
            self.poppler_path = poppler_path

        def _is_scanned_pdf(self, file_path: str) -> bool:
            """判断是否为扫描件"""
            try:
                with pdfplumber.open(file_path) as pdf:
                    sample_text = ""
                    # 检查前3页
                    for page in pdf.pages[:3]:
                        text = page.extract_text()
                        if text:
                            sample_text += text
                    # 清洗文本
                    sample_text = sample_text.strip()
                    # 如果抽取的文本极少，且有图片，则认为是扫描件
                    return len(sample_text) < 50
            except:
                return True  # 出错默认按扫描件处理
            return False

        def _process_scanned_page(self, image: Image.Image) -> str:
            """处理扫描件的单页（整页转图）"""
            return self.ocr_service.process_image(image)

        def _process_native_page(self, page) -> str:
            """处理原生PDF的单页（提取文字、表格、图片）"""
            text = ""
            # 1. 提取文字
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n\n"

            # 2. 提取表格 (pdfplumber原生)
            tables = page.extract_tables()
            for table in tables:
                text += self._convert_table_to_md(table) + "\n\n"

            # 3. 提取图片OCR
            if page.images and self.ocr_service:
                for img in page.images:
                    try:
                        # 裁剪图片区域
                        bbox = (img["x0"], img["top"], img["x1"], img["bottom"])
                        img_crop = page.crop(bbox)
                        img_obj = img_crop.to_image()
                        text += (
                            self.ocr_service.process_image(img_obj.original) + "\n\n"
                        )
                    except Exception as e:
                        print(f"PDF图片提取失败: {e}")
            return text

        def extract_text(self, file_path: str) -> str:
            """对外接口：根据类型分发"""
            if self._is_scanned_pdf(file_path):
                # 扫描件逻辑
                try:
                    full_text = ""
                    # 使用 Poppler 路径
                    pages = convert_from_path(
                        file_path, dpi=200, poppler_path=self.poppler_path
                    )
                    for i, img in enumerate(pages):
                        page_text = self._process_scanned_page(img)
                        full_text += f"[PDF第{i+1}页]\n{page_text}\n"
                    return full_text
                except Exception as e:
                    return f"扫描件转换失败: {str(e)}"
            else:
                # 原生PDF逻辑
                full_text = ""
                try:
                    with pdfplumber.open(file_path) as pdf:
                        for page in pdf.pages:
                            full_text += self._process_native_page(page)
                except Exception as e:
                    print(f"原生PDF读取失败: {e}")
                return full_text

        def _convert_table_to_md(self, table: List[List[str]]) -> str:
            """辅助方法：将pdfplumber表格转为MD"""
            if not table:
                return ""
            md = []
            # 表头
            md.append("| " + " | ".join(map(str, table[0])) + " |")
            # 分隔线
            md.append("| " + " | ".join(["---"] * len(table[0])) + " |")
            # 数据
            for row in table[1:]:
                md.append("| " + " | ".join(map(str, row)) + " |")
            return "\n".join(md)

    # ==========================================
    # 内部类 3: Word 处理器
    # ==========================================
    class _WordHandler:
        """专门处理Word逻辑"""

        def __init__(self, ocr_service: "_OcrService"):
            self.ocr_service = ocr_service

        def extract_text(self, file_path: str) -> str:
            full_text = ""
            try:
                doc = DocxDocument(file_path)
                # 1. 段落
                for para in doc.paragraphs:
                    if para.text.strip():
                        full_text += para.text + "\n\n"
                # 2. 表格转 Markdown
                for table in doc.tables:
                    table_lines = []
                    for row in table.rows:
                        cells = [cell.text.strip() for cell in row.cells]
                        table_lines.append("| " + " | ".join(cells) + " |")
                    if table_lines:
                        # 添加表头分隔线
                        if len(table_lines) > 1:
                            header_cell_count = (
                                len(table_lines[0].split("|")) - 2
                            )  # -2 for the leading/trailing empty strings
                            sep = "| " + " | ".join(["---"] * header_cell_count) + " |"
                            table_lines.insert(1, sep)
                        full_text += "\n".join(table_lines) + "\n\n"

                # 3. 图片OCR (临时目录保存图片)
                if self.ocr_service:
                    with tempfile.TemporaryDirectory() as tmpdir:
                        for i, shape in enumerate(doc.inline_shapes):
                            if shape.type == 3:  # 图片
                                try:
                                    image = shape.image
                                    img_path = os.path.join(tmpdir, f"img_{i}.png")
                                    with open(img_path, "wb") as f:
                                        f.write(image.blob)
                                    img = Image.open(img_path)
                                    ocr_result = self.ocr_service.process_image(img)
                                    full_text += ocr_result + "\n"
                                except Exception as e:
                                    full_text += f"[Word图片OCR失败: {e}]\n"
            except Exception as e:
                print(f"处理Word文档失败: {e}")
            return full_text

    # ==========================================
    # 主类的辅助方法 (路径配置)
    # ==========================================

    def _setup_tesseract_path(self):
        """智能配置 Tesseract OCR 引擎路径"""
        if platform.system() != "Windows":
            return  # Linux/Mac 依赖系统 PATH

        # 1. 尝试常见的默认路径
        default_paths = [
            r"E:\damoxing\python-project\help_ocr\tesseract\install\tesseract.exe",
            r"C:\Program Files\Tesseract-OCR\tesseract.exe",
            r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
        ]
        for path in default_paths:
            if os.path.exists(path):
                pytesseract.pytesseract.tesseract_cmd = path
                print(f"自动发现 Tesseract 路径: {path}")
                return
        print("警告: 未在 Windows 上找到 Tesseract，OCR 功能将无法使用！")

    def _get_poppler_path(self):
        """智能获取 Poppler 路径"""
        if platform.system() != "Windows":
            return None

        default_paths = [
            r"E:\damoxing\python-project\help_ocr\poppler-26.02.0\Library\bin",
            r"C:\Program Files\poppler\Library\bin",
            r".\poppler\Library\bin",
        ]
        for path in default_paths:
            exe_name = "pdftoppm.exe" if platform.system() == "Windows" else "pdftoppm"
            if os.path.exists(os.path.join(path, exe_name)):
                print(f"自动发现 Poppler 路径: {path}")
                return path
        print("警告: 未在 Windows 上找到 Poppler 路径，PDF 扫描件处理可能会失败。")
        return None

    # ==========================================
    # 主类的对外接口
    # ==========================================

    def load_document(self, file_path: str) -> List[Document]:
        ext = os.path.splitext(file_path)[1].lower()
        text = ""
        if ext == ".pdf":
            text = self.pdf_handler.extract_text(file_path)
            f_type = "pdf"
        elif ext == ".docx":
            text = self.word_handler.extract_text(file_path)
            f_type = "docx"
        elif ext == ".txt":
            # 探测编码
            with open(file_path, "rb") as f:
                raw = f.read(10240)
                result = chardet.detect(raw)
                encoding = result["encoding"] if result["confidence"] > 0.7 else "utf-8"
            try:
                with open(file_path, "r", encoding=encoding) as f:
                    text = f.read()
            except:
                with open(file_path, "r", encoding="gbk", errors="ignore") as f:
                    text = f.read()
            f_type = "txt"
        else:
            raise ValueError(f"不支持的格式: {ext}")

        if not text.strip():
            raise ValueError("文档内容为空")

        return [
            Document(
                page_content=text, metadata={"source": file_path, "file_type": f_type}
            )
        ]

    def load_directory(self, directory_path: str) -> List[Document]:
        """加载目录"""
        all_documents = []
        try:
            for root, _, files in os.walk(directory_path):
                for file_name in files:
                    file_path = os.path.join(root, file_name)
                    ext = os.path.splitext(file_path)[1].lower()
                    if ext not in [".txt", ".pdf", ".docx"]:
                        continue
                    docs = self.load_document(file_path)
                    all_documents.extend(docs)
            return self.split_documents(all_documents)
        except Exception as e:
            print(f"加载目录失败: {e}")
            return []

    def split_documents(self, documents: List[Document]) -> List[Document]:
        """切分文档"""
        if self.text_splitter is None:
            raise ValueError("切分器未初始化")
        split_docs = self.text_splitter.split_documents(documents)
        # 保存切块到本地
        self.save_chunks_to_local(split_docs)
        # 设置 Chunk ID
        for doc in split_docs:
            doc.metadata["chunk_id"] = self._global_chunk_counter
            self._global_chunk_counter += 1
        return split_docs

    def save_chunks_to_local(
        self, documents: List[Document], output_dir: str = "data/chunk_documents"
    ) -> None:
        """保存切块到本地"""
        if not documents:
            print("警告: 没有文档块需要保存。")
            return
        os.makedirs(output_dir, exist_ok=True)
        for doc in documents:
            chunk_id = doc.metadata.get("chunk_id", 0)
            source_path = doc.metadata.get("source", "unknown_file")
            base_name = os.path.splitext(os.path.basename(source_path))[0]
            filename = f"chunk_{chunk_id:04d}_{base_name}.md"
            filepath = os.path.join(output_dir, filename)

            # 构建内容 (带元数据头部)
            metadata_str = "---\n"
            for key, value in doc.metadata.items():
                value_str = str(value).replace("\n", " ")
                metadata_str += f"{key}: {value_str}\n"
            metadata_str += "---\n\n"
            content = metadata_str + doc.page_content

            try:
                with open(filepath, "w", encoding="utf-8") as f:
                    f.write(content)
            except Exception as e:
                print(f"保存文件失败 {filepath}: {e}")

    def process_documents(self, input_path: str) -> List[Document]:
        """统一入口"""
        if os.path.isfile(input_path):
            return self.load_document(input_path)
        elif os.path.isdir(input_path):
            return self.load_directory(input_path)
        else:
            raise ValueError(f"无效路径: {input_path}")

# 全局实例
documentProcessor = DocumentProcessor()
